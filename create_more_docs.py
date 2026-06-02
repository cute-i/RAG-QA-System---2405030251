import os
from docx import Document


def create_docx_document(file_path, title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    
    for para in content:
        doc.add_paragraph(para)
    
    doc.save(file_path)
    print(f"Created: {file_path}")


def main():
    os.makedirs("documents", exist_ok=True)
    
    docs = [
        {
            "filename": "命名实体识别技术.docx",
            "title": "命名实体识别技术",
            "content": [
                "命名实体识别（Named Entity Recognition，NER）是自然语言处理中的一项重要任务，其目标是识别文本中具有特定意义的实体。",
                "常见的实体类型包括：人名（PER）、地名（LOC）、机构名（ORG）、时间（TIME）、日期（DATE）、数字（NUM）等。",
                "传统的NER方法包括基于规则的方法、基于统计的方法（如隐马尔可夫模型HMM、条件随机场CRF等）。",
                "基于深度学习的NER方法包括：双向LSTM+CRF、BERT+CRF等，这些方法在准确率上取得了显著提升。",
                "NER在信息抽取、问答系统、机器翻译等领域都有广泛应用。"
            ]
        },
        {
            "filename": "情感分析技术.docx",
            "title": "情感分析技术",
            "content": [
                "情感分析（Sentiment Analysis）是指从文本中识别和提取情感倾向的技术，包括正面、负面、中性等。",
                "情感分析的应用场景包括：社交媒体监控、产品评论分析、舆情分析、客户反馈分析等。",
                "情感分析的方法包括：基于词典的方法、机器学习方法（如朴素贝叶斯、支持向量机）、深度学习方法（如CNN、RNN、Transformer）。",
                "情感分析的挑战包括：讽刺和反语的识别、上下文依赖、多模态情感分析等。",
                "常用的情感分析数据集包括IMDB电影评论、Amazon产品评论、Twitter情感数据集等。"
            ]
        },
        {
            "filename": "机器翻译技术.docx",
            "title": "机器翻译技术",
            "content": [
                "机器翻译（Machine Translation，MT）是指利用计算机将一种自然语言转换为另一种自然语言的技术。",
                "机器翻译的发展历程包括：基于规则的机器翻译、基于统计的机器翻译、基于神经网络的机器翻译。",
                "神经机器翻译（Neural Machine Translation，NMT）是当前最先进的机器翻译方法，基于Encoder-Decoder架构。",
                "Transformer架构的提出进一步提升了机器翻译的质量，Google的Transformer模型和OpenAI的GPT系列都采用了这种架构。",
                "机器翻译的评估指标包括BLEU、ROUGE、METEOR等，其中BLEU是最常用的自动评估指标。"
            ]
        },
        {
            "filename": "文本摘要技术.docx",
            "title": "文本摘要技术",
            "content": [
                "文本摘要（Text Summarization）是指从文本中提取关键信息并生成简洁摘要的技术。",
                "文本摘要分为两种类型：抽取式摘要（Extractive）和生成式摘要（Abstractive）。",
                "抽取式摘要从原文中选取重要句子组成摘要，而生成式摘要则生成新的句子来表达原文的核心内容。",
                "常用的抽取式方法包括：TF-IDF、TextRank、基于深度学习的方法等。",
                "生成式摘要方法包括：Seq2Seq模型、Transformer模型、预训练语言模型（如BERT、GPT）等。"
            ]
        },
        {
            "filename": "问答系统技术.docx",
            "title": "问答系统技术",
            "content": [
                "问答系统（Question Answering System）是指能够理解自然语言问题并给出准确回答的系统。",
                "问答系统分为多种类型：事实型问答、开放域问答、阅读理解式问答、对话式问答等。",
                "问答系统的核心技术包括：问题理解、文档检索、答案抽取、答案生成等。",
                "基于深度学习的问答系统方法包括：基于BERT的阅读理解模型、基于Transformer的问答模型等。",
                "问答系统的应用包括：智能客服、知识库问答、教育辅助、医疗咨询等领域。"
            ]
        },
        {
            "filename": "词嵌入技术详解.docx",
            "title": "词嵌入技术详解",
            "content": [
                "词嵌入（Word Embedding）是将离散的词语转换为连续向量表示的技术，能够捕捉词语的语义信息。",
                "Word2Vec是最早的词嵌入模型之一，包括CBOW（Continuous Bag of Words）和Skip-gram两种架构。",
                "GloVe（Global Vectors for Word Representation）通过全局词频统计来学习词向量。",
                "FastText在Word2Vec的基础上引入了字符级别的n-gram特征，能够更好地处理未登录词。",
                "ELMo（Embeddings from Language Models）使用预训练的双向LSTM模型生成动态词向量。",
                "BERT（Bidirectional Encoder Representations from Transformers）使用Transformer架构，能够捕捉上下文相关的词表示。"
            ]
        },
        {
            "filename": "预训练语言模型.docx",
            "title": "预训练语言模型",
            "content": [
                "预训练语言模型（Pre-trained Language Models）是在大规模语料上预先训练的语言模型，能够学习通用的语言知识。",
                "ELMo是第一个成功应用预训练的模型，使用双向LSTM在大规模文本上进行预训练。",
                "GPT（Generative Pre-trained Transformer）是OpenAI提出的生成式预训练模型，基于Transformer解码器。",
                "BERT（Bidirectional Encoder Representations from Transformers）是Google提出的基于Transformer编码器的预训练模型。",
                "GPT-2、GPT-3、GPT-4是GPT系列的后续版本，参数规模不断增大，能力不断提升。",
                "预训练语言模型的微调（Fine-tuning）是指在特定任务上进行训练，使其适应具体的应用场景。"
            ]
        },
        {
            "filename": "NLP数据集与评估.docx",
            "title": "NLP数据集与评估",
            "content": [
                "NLP任务的评估需要使用标准数据集，常见的数据集包括：",
                "GLUE（General Language Understanding Evaluation）是一个包含多个NLP任务的基准数据集。",
                "SQuAD（Stanford Question Answering Dataset）是一个阅读理解数据集，包含问题和对应的段落。",
                "CoNLL-2003是一个命名实体识别数据集，包含英语和德语的标注数据。",
                "IMDB电影评论数据集用于情感分析任务，包含50000条电影评论。",
                "常用的评估指标包括：准确率（Accuracy）、精确率（Precision）、召回率（Recall）、F1分数、BLEU分数等。",
                "对于生成任务，人工评估也是重要的评估方式，能够评估生成文本的质量和流畅度。"
            ]
        }
    ]
    
    for doc in docs:
        filepath = os.path.join("documents", doc["filename"])
        create_docx_document(filepath, doc["title"], doc["content"])
    
    print("\n✅ 所有新文档已创建完成！")
    print(f"总共创建了 {len(docs)} 个新文档")


if __name__ == "__main__":
    main()
