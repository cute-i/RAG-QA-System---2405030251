import os
from docx import Document


def create_docx_document(file_path, title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    
    for para in content:
        doc.add_paragraph(para)
    
    doc.save(file_path)
    print(f"Created: {file_path}")


def main():
    os.makedirs("documents", exist_ok=True)
    
    docs = [
        {
            "filename": "自然语言处理基础.docx",
            "title": "自然语言处理基础",
            "content": [
                "自然语言处理（Natural Language Processing，NLP）是人工智能领域的一个重要分支，它致力于使计算机能够理解、解释和生成人类语言。",
                "NLP的主要任务包括：文本分类、情感分析、命名实体识别、机器翻译、问答系统等。这些任务涵盖了从词法分析到语义理解的各个层面。",
                "近年来，深度学习技术的发展极大地推动了NLP的进步。特别是Transformer架构的提出，为NLP带来了革命性的变化。",
                "词嵌入（Word Embedding）是NLP中的核心技术之一，它将离散的词语转换为连续的向量表示，使得计算机能够更好地理解词语之间的语义关系。",
                "常用的词嵌入模型包括Word2Vec、GloVe和BERT等。这些模型能够捕捉词语的语义信息，为后续的NLP任务提供有力支持。"
            ]
        },
        {
            "filename": "Transformer架构详解.docx",
            "title": "Transformer架构详解",
            "content": [
                "Transformer是一种基于自注意力机制的深度学习架构，由Vaswani等人在2017年提出。它彻底改变了序列建模的方式。",
                "Transformer主要由编码器（Encoder）和解码器（Decoder）两部分组成。编码器负责处理输入序列，解码器负责生成输出序列。",
                "自注意力机制（Self-Attention）是Transformer的核心创新。它允许模型在处理每个位置时，关注输入序列中的所有位置。",
                "多头注意力（Multi-Head Attention）机制通过并行计算多个注意力头，能够捕捉不同类型的依赖关系，增强模型的表达能力。",
                "位置编码（Positional Encoding）为模型提供了序列的顺序信息，弥补了自注意力机制本身不具备顺序感知能力的缺陷。"
            ]
        },
        {
            "filename": "文本分类方法.docx",
            "title": "文本分类方法",
            "content": [
                "文本分类是NLP中的一项基础任务，其目标是将文本自动归类到预定义的类别中。",
                "传统的文本分类方法包括：基于规则的方法、基于统计的方法（如朴素贝叶斯、支持向量机等）。",
                "基于深度学习的文本分类方法包括：卷积神经网络（CNN）、循环神经网络（RNN）、Transformer等。",
                "文本分类的一般流程包括：数据预处理、特征提取、模型训练和预测。",
                "评估指标通常包括准确率（Accuracy）、精确率（Precision）、召回率（Recall）和F1分数等。"
            ]
        },
        {
            "filename": "注意力机制原理.docx",
            "title": "注意力机制原理",
            "content": [
                "注意力机制是一种模拟人类注意力过程的计算机制，它允许模型在处理信息时有选择地关注输入的不同部分。",
                "注意力机制的核心思想是：根据查询向量（Query）、键向量（Key）和值向量（Value）计算注意力权重。",
                "缩放点积注意力（Scaled Dot-Product Attention）是最常用的注意力计算方式，其计算公式为：Attention(Q,K,V) = softmax(QK^T / sqrt(d_k))V。",
                "注意力机制在机器翻译、文本摘要、问答系统等多个NLP任务中都取得了显著的效果。",
                "除了自注意力机制，还有交叉注意力（Cross-Attention）等变体，用于处理不同序列之间的注意力关系。"
            ]
        },
        {
            "filename": "深度学习在NLP中的应用.docx",
            "title": "深度学习在NLP中的应用",
            "content": [
                "深度学习技术的快速发展为NLP带来了前所未有的机遇。深度学习模型能够自动学习文本的复杂特征表示。",
                "预训练语言模型（如BERT、GPT等）通过在大规模语料上进行预训练，学习到通用的语言知识，然后通过微调适应特定任务。",
                "BERT（Bidirectional Encoder Representations from Transformers）是一种基于Transformer的预训练模型，能够捕捉上下文信息。",
                "GPT（Generative Pre-trained Transformer）是一种生成式预训练模型，在文本生成任务中表现出色。",
                "深度学习在NLP中的应用包括：情感分析、机器翻译、文本生成、问答系统、对话系统等多个领域。"
            ]
        }
    ]
    
    for doc in docs:
        filepath = os.path.join("documents", doc["filename"])
        create_docx_document(filepath, doc["title"], doc["content"])
    
    print("\n✅ 所有示例文档已创建完成！")


if __name__ == "__main__":
    main()
